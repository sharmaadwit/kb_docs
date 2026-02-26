import docx
from docx.shared import Pt
import sys
import re

def markdown_to_docx(md_path, docx_path):
    doc = docx.Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    table_data = []
    in_table = False
    
    for line in lines:
        line = line.strip()
        
        # Handle Tables
        if '|' in line and '---' not in line:
            in_table = True
            row = [cell.strip() for cell in line.split('|') if cell.strip()]
            if row:
                table_data.append(row)
            continue
        elif in_table:
            # Create Table in Docx
            if table_data:
                table = doc.add_table(rows=0, cols=len(table_data[0]))
                table.style = 'Table Grid'
                for row_data in table_data:
                    row_cells = table.add_row().cells
                    for i, text in enumerate(row_data):
                        row_cells[i].text = text
            table_data = []
            in_table = False
            if not line: continue

        # Handle Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        # Handle Bullet Points
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        # Handle Body Text
        elif line:
            # Basic bolding check
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            doc.add_paragraph()

    # Final table check if file ends with a table
    if in_table and table_data:
        table = doc.add_table(rows=0, cols=len(table_data[0]))
        table.style = 'Table Grid'
        for row_data in table_data:
            row_cells = table.add_row().cells
            for i, text in enumerate(row_data):
                row_cells[i].text = text

    doc.save(docx_path)
    print(f"Successfully saved to {docx_path}")

if __name__ == "__main__":
    markdown_to_docx(sys.argv[1], sys.argv[2])
